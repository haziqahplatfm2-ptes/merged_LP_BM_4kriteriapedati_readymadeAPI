import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from io import BytesIO

# --- 1. KONFIGURASI ---
genai.configure(api_key=st.secrets["GOOGLE_API_KEY"])

@st.cache_resource
def find_working_model():
    try:
        for m in genai.list_models():
            if 'generateContent' in m.supported_generation_methods:
                return m.name
    except:
        return "models/gemini-1.5-flash"
    return "models/gemini-1.5-flash"

selected_model_name = find_working_model()
model = genai.GenerativeModel(selected_model_name)

# --- 2. LOGIK AI (KRITERIA GABUNGAN DALAM BAHASA MELAYU) ---
def generate_advanced_plan_bm(topic, syllabus, extra_context):
    prompt = f"""
    Topik: {topic}. Kod Sukatan: {syllabus}. Konteks Tambahan: {extra_context}.
    Hasilkan satu rancangan pengajaran profesional dalam BAHASA MELAYU sepenuhnya. 
    Sila elakkan penggunaan sebarang istilah Bahasa Inggeris di seluruh kandungan dokumen.
    
    PERATURAN UTAMA PENGGUNAAN ISTILAH:
    1. JANGAN SEKALI-KALI menggunakan perkataan 'MURID'. Gantikan KESEMUANYA dengan istilah 'PELAJAR' di seluruh rancangan pengajaran tanpa sebarang pengecualian.

    PERATURAN KRITIKAL FORMAT KANDUNGAN:
    1. JANGAN gunakan simbol double asterisk (**) di mana-mana bahagian respons.
    2. JANGAN gunakan senarai bentuk bullet (seperti -, *, •). Anda WAJIB menggunakan senarai bernombor (1, 2, 3...) secara eksklusif untuk semua bahagian bersenarai.
    3. Semua penanda bahagian (SECTION:) di bawah mestilah ditulis dalam HURUF BESAR sepenuhnya.

    Gunakan penanda bahagian (SECTION:) yang TEPAT ini untuk struktur dokumen:
    
    SECTION: TOPIK
    {topic}
    
    SECTION: OBJEKTIF PEMBELAJARAN
    [Sediakan tepat 4 isi menggunakan format bernombor 1., 2., 3., 4.]
    
    SECTION: HASIL PEMBELAJARAN
    [Sediakan tepat 4 isi menggunakan format bernombor 1., 2., 3., 4.]
    
    SECTION: KRITERIA KEJAYAAN
    [Sediakan tepat 4 isi menggunakan format bernombor 1., 2., 3., 4.]
    
    SECTION: PRASYARAT
    [Sediakan 1 pernyataan tentang pengetahuan sedia ada pelajar]
    
    SECTION: KATA KUNCI
    [Sediakan 6 item kata kunci penting yang dipisahkan oleh tanda koma sahaja. Jangan buat bentuk senarai.]
    
    SECTION: HOTS
    1. Menganalisis: [Sediakan butiran ringkas berasaskan topik]
    2. Menilai: [Sediakan butiran ringkas berasaskan topik]
    3. Mencipta: [Sediakan butiran ringkas berasaskan topik]
    4. Mengaplikasi: [Sediakan butiran ringkas berasaskan topik]
    
    SECTION: KEWARGANEGARAAN DIGITAL
    [Sediakan tepat 4 isi bernombor 1., 2., 3., 4. tentang etika penggunaan teknologi/Chromebook/Canva/YouTube oleh PELAJAR]

    SECTION: KANDUNGAN PEMBUKAAN LESEN
    [Aktiviti set induksi/pencetus minat dan pelan transisi sebelum pembelajaran bermula]

    SECTION: STRATEGI DIFERENSIASI (HIJAU)
    1. HA (Higher Achiever): [1 aktiviti mencabar untuk pelajar berpencapaian tinggi]

    SECTION: STRATEGI DIFERENSIASI (KUNING)
    1. MA (Medium Achiever): [1 aktiviti teras untuk pelajar berpencapaian sederhana]

    SECTION: STRATEGI DIFERENSIASI (MERAH)
    1. LA (Lower Achiever): [1 aktiviti bersofisikasi/berpandu untuk pelajar berpencapaian rendah]

    SECTION: AKTIVITI PEMBELAJARAN TERADUN SATU (15 MINIT)
    1. Aktiviti 1: [Penerangan aktiviti teradun pertama]
    2. Persediaan Pensyarah: [Langkah demi langkah sebelum sesi bermula]
    3. Objektif: [3 isi bernombor]
    4. Tugas Pelajar: [Butiran langkah demi langkah khusus untuk PELAJAR]

    SECTION: AKTIVITI PEMBELAJARAN TERADUN DUA (15 MINIT)
    1. Aktiviti 2: [Penerangan aktiviti teradun kedua]
    2. Persediaan Pensyarah: [Langkah demi langkah sebelum sesi bermula]
    3. Objektif: [3 isi bernombor]
    4. Tugas Pelajar: [Butiran langkah demi langkah khusus untuk PELAJAR]

    SECTION: GRID ALIRAN PEDATI
    BLOCK_START: P: PERSEDIAAN (BELAJAR)
    PENSYARAH: [Langkah tindakan yang sejajar dengan topik]
    PELAJAR: [Tugasan tindakan/kerja chromebook yang sejajar dengan topik]
    BLOCK_END
    
    BLOCK_START: E: PENGLIBATAN (TEROKA)
    PENSYARAH: [Langkah tindakan yang sejajar dengan topik]
    PELAJAR: [Tugasan tindakan/kerja chromebook yang sejajar dengan topik]
    BLOCK_END

    BLOCK_START: D.A: PENYAMPAIAN DAN APLIKASI
    PENSYARAH: [Langkah tindakan yang sejajar dengan topik]
    PELAJAR: [Tugasan tindakan/kerja chromebook yang sejajar dengan topik]
    BLOCK_END

    BLOCK_START: T.I: UJIAN DAN PENILAIAN
    PENSYARAH: [Langkah tindakan yang sejajar dengan topik]
    PELAJAR: [Tugasan tindakan/kerja chromebook yang sejajar dengan topik]
    BLOCK_END
    
    SECTION: PLENARI (TIKET KELUAR)
    [Aktiviti penutup kelas sekitar 2-3 minit untuk menguji kefahaman pelajar]

    SECTION: KERJA RUMAH
    [Tugasan latihan susulan berdasarkan topik pelajaran]

    SECTION: CADANGAN TUGASAN MELANGKAH KE HADAPAN
    [Aktiviti pencetus minat dan pelan transisi untuk sesi pengajaran hari esok]
    """
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"System Error: {str(e)}"

def add_page_number(run):
    """Memasukkan nombor halaman dinamik Word di bahagian atas tengah."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# --- 3. LOGIK EKSPORT WORD ---
def create_word_export(topic, syllabus, text):
    doc = Document()
    
    # Peraturan 1: Saiz kertas LETTER dengan margin 0.5 inci di setiap penjuru
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # Peraturan 7: Nombor halaman di bahagian atas tengah
        header = section.header
        header_p = header.paragraphs[0]
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_run = header_p.add_run()
        header_run.font.name = 'Arial'
        header_run.font.size = Pt(10)
        add_page_number(header_run)

    # Peraturan 2 & 8: Tajuk Utama dalam HURUF BESAR (Saiz Font 14)
    main_title = f'PTES RANCANGAN PENGAJARAN GABUNGAN UNIVERSAL & PEDATI: {topic}'.upper()
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(main_title)
    title_run.font.size = Pt(14)
    title_run.bold = True
    
    # Peraturan 3 & 8: Konfigurasi perenggan normal (Arial 12, Jarak baris TUNGGAL)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.0
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # Jadual Pentadbiran (Admin Table)
    admin_table = doc.add_table(rows=3, cols=4)
    admin_table.style = 'Table Grid'
    labels = [["Minggu Ke:", "Tarikh:"], ["Bilangan Pelajar:", "Hari:"], ["Tempat / Makmal:", "Durasi (minit):"]]
    for r in range(3):
        admin_table.cell(r, 0).text = labels[r][0]
        admin_table.cell(r, 2).text = labels[r][1]
        
    for row in admin_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.size = Pt(12)
    doc.add_paragraph().paragraph_format.line_spacing = 1.0

    # Proses Pembahagian Blok Kandungan Utama
    sections = text.split('SECTION:')
 
    for section in sections:
        if not section.strip(): continue
        lines = section.strip().split('\n')
        
        # Peraturan 2 & 4: Paksa huruf besar untuk semua tajuk bahagian & buang asterisk
        title = lines[0].strip().replace("**", "").upper()
        body_content = "\n".join(lines[1:]).strip()
        
        doc_heading = doc.add_paragraph()
        doc_heading.paragraph_format.line_spacing = 1.0
        h_run = doc_heading.add_run(title)
        h_run.bold = True
        h_run.font.size = Pt(14)  # Peraturan 8: Saiz font tajuk 14

        # Peraturan 6: Pemprosesan jadual matriks kotak KATA KUNCI rata tengah
        if "KATA KUNCI" in title or "KEYWORDS" in title:
            raw_keywords_text = " ".join([l.strip() for l in lines[1:] if l.strip()])
            keyword_items = [kw.strip() for kw in raw_keywords_text.split(",") if kw.strip()]
            
            kw_table = doc.add_table(rows=2, cols=3)
            kw_table.style = 'Table Grid'
            
            idx = 0
            for r in range(2):
                for c in range(3):
                    if idx < len(keyword_items):
                        cell = kw_table.cell(r, c)
                        cell.text = keyword_items[idx]
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.paragraph_format.line_spacing = 1.0
                        if p.runs:
                            p.runs[0].font.size = Pt(12)
                        idx += 1
            doc.add_paragraph().paragraph_format.line_spacing = 1.0
            
        elif "GRID ALIRAN PEDATI" in title or "PEDATI FLOW GRID" in title:
            blocks = body_content.split("BLOCK_START:")
            for block in blocks:
                if not block.strip(): continue
                block_data = block.split("BLOCK_END")[0].strip().split('\n')
                
                heading_title = block_data[0].strip().replace("**", "").upper() # Peraturan 2: Huruf besar tajuk fasa aliran
                lecturer_text = ""
                pelajar_text = ""
                
                for line in block_data:
                    if line.upper().startswith("PENSYARAH:"):
                        lecturer_text = line.split(":", 1)[1].strip().replace("**", "")
                    elif line.upper().startswith("PELAJAR:"):
                        pelajar_text = line.split(":", 1)[1].strip().replace("**", "")
                
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.space_before = Pt(6)
                run = p.add_run(heading_title)
                run.bold = True
                run.font.size = Pt(12)
                
                table = doc.add_table(rows=2, cols=2)
                table.style = 'Table Grid'
                
                for row in table.rows:
                    row.cells[0].width = Inches(3.75)
                    row.cells[1].width = Inches(3.75)
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Pensyarah"
                hdr_cells[1].text = "Pelajar"
                
                data_cells = table.rows[1].cells
                data_cells[0].text = lecturer_text
                data_cells[1].text = pelajar_text
                
                # Penguatkuasaan format teks dalam jadual (Arial 12, Jarak baris tunggal)
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.line_spacing = 1.0
                            for run in paragraph.runs:
                                run.font.size = Pt(12)
                doc.add_paragraph().paragraph_format.line_spacing = 1.0
        else:
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            
            content = body_content.replace("**", "")
            table.cell(0, 0).text = content
            
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.line_spacing = 1.0
                        for run in paragraph.runs:
                            run.font.size = Pt(12)  # Peraturan 8: Font teks isi saiz 12
            doc.add_paragraph().paragraph_format.line_spacing = 1.0
     
    # Jadual Kelulusan HOD (Ketua Jabatan)
    doc.add_page_break()
    
    hod_heading = doc.add_paragraph()
    hod_heading.paragraph_format.line_spacing = 1.0
    hod_run = hod_heading.add_run("ULASAN & PENGESAHAN HOD")
    hod_run.bold = True
    hod_run.font.size = Pt(14)
    
    hod_table = doc.add_table(rows=2, cols=2)
    hod_table.style = 'Table Grid'
    hod_table.cell(0, 0).text = "Ulasan / Catatan:"
    hod_table.rows[1].height = Pt(50)
    hod_table.cell(1, 0).text = "Tarikh:"
    hod_table.cell(1, 1).text = "Tandatangan / Cop Rasmi:"
    
    for row in hod_table.rows:
        for cell in row.cells:
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            if p.runs:
                p.runs[0].font.size = Pt(12)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- 4. GRAPHICAL USER INTERFACE (GUI) ---
st.set_page_config(page_title="Penjana Gabungan Universal & PEDATI", layout="wide")

st.title("🎓 PTES Rancangan Mengajar Universal & PEDATI")
st.info("Masukkan topik pelajaran, kod sukatan subjek dan maklumat tambahan untuk menjana rancangan pengajaran hibrid versi Bahasa Melayu murni.")

c1, c2 = st.columns(2)
with c1: u_topic = st.text_input("Topik Pelajaran:")
with c2: u_syllabus = st.text_input("Kod Sukatan Pelajaran:")
u_extra = st.text_area("Konteks Tambahan / Pilihan (Opsional):")

if st.button("🚀 JANA RANCANGAN PENGAJARAN LENGKAP"):
    if u_topic and u_syllabus:
        with st.spinner("AI sedang membina rancangan pengajaran berstruktur hibrid..."):
            result = generate_advanced_plan_bm(u_topic, u_syllabus, u_extra)
            st.session_state['hybrid_plan_out_bm'] = result.replace("**", "")
    else:
        st.warning("Sila isi ruang Topik dan Kod Sukatan Pelajaran.")

if 'hybrid_plan_out_bm' in st.session_state:
    st.divider()
    st.subheader("📝 Pratonton Draf AI (Format Melayu Bersih)")
    st.text_area("Kandungan", st.session_state['hybrid_plan_out_bm'], height=400)
    doc_file = create_word_export(u_topic, u_syllabus, st.session_state['hybrid_plan_out_bm'])
    st.download_button("📥 Muat Turun ke Versi Word (.docx)", doc_file, f"RMH_Merge_{u_topic}.docx")

st.markdown("---")
st.caption("Lesson Planner Hybrid v4.0 | Pencipta: Hjh Nurul Haziqah Hj Nordin | © 2026 PTES Academic Innovation")
